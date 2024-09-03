import os
import random
import string
import platform
import subprocess
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from config.settings import MEDIA_ROOT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def set_row_height(row, height_in_inches):
    """Set the height of a row."""
    tr = row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_in_inches * 1440)))  # 1440 twips = 1 inch
    trPr.append(trHeight)

def set_cell_width(cell, width_in_inches):
    """Set the width of a cell using standard python-docx features."""
    # Set the width of the cell
    cell.width = Inches(width_in_inches)
    
    # Apply width to cell using the table cell properties (tcPr) for better compatibility
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_in_inches * 1440)))  # 1440 twips = 1 inch
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def adjust_specific_cell_width(table, cell_name, new_width_in_inches):
    """Adjust the width of a specific cell in the table by searching for the cell's text."""
    cell_found = False
    for cell in table.rows[0].cells:
        if cell.text.strip() == cell_name:
            set_cell_width(cell, new_width_in_inches)
            cell_found = True
            break

    if not cell_found:
        print(f"Cell '{cell_name}' not found in the table.")

def make_cell_bold(cell):
    """Make the content of a cell bold."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

def format_additional_data(additional_data):
    s_no = []
    description = []
    qty = []
    uom = []
    item_rate = []
    amount = []
    disc = []
    taxable = []

    for data in additional_data:
        s_no.append(data[0])
        description.append(data[1])
        qty.append(data[2])
        uom.append(data[3])
        item_rate.append(data[4])
        amount.append(data[5])
        disc.append(data[6])
        taxable.append(data[7])

    # Joining the lists with newline characters
    s_no_str = '\n'.join(s_no)
    description_str = '\n'.join(description)
    qty_str = '\n'.join(qty)
    uom_str = '\n'.join(uom)
    item_rate_str = '\n'.join(item_rate)
    amount_str = '\n'.join(amount)
    disc_str = '\n'.join(disc)
    taxable_str = '\n'.join(taxable)

    return s_no_str, description_str, qty_str, uom_str, item_rate_str, amount_str, disc_str, taxable_str

def create_sales_order_doc(product_data, Cust_data):
    doc = Document()

    # Set page orientation to landscape
    section = doc.sections[0]
    section.orientation = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_width, new_height = section.page_height, section.page_width
    section.page_width = Inches(14.54)
    section.page_height = Inches(11.69)

    header_text = Cust_data.get('{{Doc Header}}')
    heading = doc.add_heading(header_text, level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.runs[0]
    run.font.size = Pt(24)
    run.font.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_before = Pt(0)  # Reduce space before the heading
    heading.paragraph_format.space_after = Pt(0)  # Set the space after the first heading to 0
   
    doc.add_paragraph('')
    doc.add_paragraph('')

    # Add "BILL OF SUPPLY" below the header if the header is "SALES QUOTATION"
    if header_text == "SALES QUOTATION":
        subheading = doc.add_paragraph("BILL OF SUPPLY")
        subheading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subheading_run = subheading.runs[0]
        subheading_run.bold = True
        subheading_run.font.size = Pt(12)
        subheading_run.font.color.rgb = RGBColor(0, 0, 0)

    def create_table_with_columns(num_columns, table_name, cell_names=None, bg_color=None, specific_row_height=None):
        table = doc.add_table(rows=1, cols=num_columns)
        table.style = 'Table Grid'
        table.autofit = False
        column_width = Inches(6.0 / num_columns)
        for i in range(num_columns):
            cell = table.cell(0, i)
            cell.width = column_width
            cell.text = cell_names[i] if cell_names and i < len(cell_names) else f"{table_name} Column {i + 1}"
            if bg_color:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), bg_color)
                cell._element.get_or_add_tcPr().append(shading_elm)
        if specific_row_height:
            set_row_height(table.rows[0], specific_row_height)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return table

    # Example usage:
    s_no_str, description_str, qty_str, uom_str, item_rate_str, amount_str, disc_str, taxable_str = format_additional_data(product_data)
    t1 = create_table_with_columns(3, "t1", cell_names=["Customer Billing Detail", "{{number_lbl}} {{number_value}}", "{{date_lbl}} {{date_value}}"], bg_color="87CEFA")
    t2 = create_table_with_columns(2, "t2", cell_names=["{{cust_name}}"'\n' "{{address}}  {{country}}", "Phone : {{phone}}" '\n' "Destination : {{dest}}"], specific_row_height=0.5)
    t3 = create_table_with_columns(8, "t3", cell_names=["S No", "Description", "Qty", "UOM", "Item Rate", "Amount", "Disc", "Taxable"], bg_color="87CEFA")
    t4 = create_table_with_columns(8, "t4", cell_names=[s_no_str, description_str, qty_str, uom_str, item_rate_str, amount_str, disc_str, taxable_str], specific_row_height=0.9)
    t5 = create_table_with_columns(8, "t5", cell_names=["", "Total", "{{qty_ttl}}", "", "", "{{amt_ttl}}", "", "{{txbl_ttl}}"], bg_color="87CEFA")
    t6 = create_table_with_columns(2, "t6", cell_names=['\n' "Bill Amount In Words : {{Bill_amt_wd}}" '\n''\n' "Tax Amount In Words : {{Tax_amt_wd}}" '\n''\n' "Remark : {{number_value}}" '\n', "Sub Total:"'                                                            '"{{amt_ttl}}" '\n' "Discount Amt:"'                                                    '"{{discount}}" '\n' "Round Off:"'                                                            '"{{Rnd_off}}" '\n' "Bill Total:"'                                                             '"{{txbl_ttl}}" '\n' "------------------------------------------------------------------------------------" '\n' "Party Old Balance:"'                                            '"{{Party_Old_B}}" '\n'  "------------------------------------------------------------------------------------" '\n' "{{Net_lbl}}:    "'                                                     '"{{Net_Ttl}}"'\n' ], specific_row_height=0.5)
    t7 = create_table_with_columns(1, "t7", cell_names=["\n" "Declaration: " '\n' "We declare that this invoice shows the actual price of the goods/services described and"  '\n' "that all particulars are true and correct."'\n' 'Original For Recipient' '\t''\t''\t''\t''\t''\t''\t''\t''\t''\t''\t''\t''\t' "Authorised Signatory" '\n'], specific_row_height=0.5)

    # adjust_specific_cell_width(t1, "Customer Billing Detail", 2)
    adjust_specific_cell_width(t1, "{{number_lbl}} {{number_value}}", 3)
    adjust_specific_cell_width(t1, "{{date_lbl}} {{date_value}}", 1)
    

    # Adjust the width of cells in table t3
    adjust_specific_cell_width(t3, "S No", 0.1)
    adjust_specific_cell_width(t3, "Description", 2.5)
    adjust_specific_cell_width(t3, "Qty", 0.2)
    adjust_specific_cell_width(t3, "UOM", 0.3)
    adjust_specific_cell_width(t3, "Item Rate", 1.1)
    adjust_specific_cell_width(t3, "Amount", 1.4)
    adjust_specific_cell_width(t3, "Disc", 1.0)
    adjust_specific_cell_width(t3, "Taxable", 1.0)

    replacements = Cust_data
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)

     # Make content in t1, t3, and t5 bold
    for table in [t1, t3, t5, t6]:
        for row in table.rows:
            for cell in row.cells:
                make_cell_bold(cell)

    unique_code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))

     # Save to specific directory inside MEDIA_ROOT
    output_dir = os.path.join(MEDIA_ROOT, 'sales order receipt')
    os.makedirs(output_dir, exist_ok=True)  # Ensure directory exists
    
    if header_text == "SALES QUOTATION":
        final_file_path = os.path.join(output_dir, f"{replacements.get('{{cust_name}}')}_sales_Invoice_receipt_{unique_code}.docx")
    else:  
        final_file_path = os.path.join(output_dir, f"{replacements.get('{{cust_name}}')}_sales_Order_receipt_{unique_code}.docx")
    doc.save(final_file_path)

    # Convert the Word document to PDF and return the PDF file path
    pdf_file_path = convert_docx_to_pdf(final_file_path)
    os.remove(final_file_path)
    return pdf_file_path


    # """Convert a Word document to PDF using win32com.client."""
    # # Initialize the COM library
    # pythoncom.CoInitialize()
    
    # try:
    #     word = win32.Dispatch('Word.Application')

    #     # Use absolute path to ensure Word can find the file
    #     docx_full_path = os.path.abspath(docx_file_path)

    #     doc = word.Documents.Open(docx_full_path)

    #     # Generate unique PDF file path to avoid overwriting
    #     pdf_file_path = os.path.splitext(docx_full_path)[0] + ".pdf"
    #     pdf_counter = 1
    #     while os.path.exists(pdf_file_path):
    #         pdf_file_path = os.path.splitext(docx_full_path)[0] + f"_{pdf_counter}.pdf"
    #         pdf_counter += 1

    #     # Save the document as a PDF
    #     doc.SaveAs(os.path.abspath(pdf_file_path), FileFormat=17)
    #     doc.Close()
    #     word.Quit()

    # finally:
    #     # Uninitialize the COM library
    #     pythoncom.CoUninitialize()



def convert_docx_to_pdf(docx_file_path):
    """Convert a Word document to PDF using platform-specific methods."""
    
    # Ensure the DOCX file exists
    if not os.path.exists(docx_file_path):
        raise FileNotFoundError(f"The file {docx_file_path} does not exist.")
    
    # Define the output directory where the PDF should be saved
    output_dir = os.path.join(MEDIA_ROOT, 'sales order receipt')

    # Ensure the output directory exists; if not, create it
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate the PDF file path in the output directory
    pdf_file_name = os.path.splitext(os.path.basename(docx_file_path))[0] + ".pdf"
    pdf_file_path = os.path.join(output_dir, pdf_file_name)
    
    # Detect the operating system
    current_os = platform.system()
    
    if current_os == 'Windows':
        # On Windows, use LibreOffice (ensure path is correct)
        libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"  # Update if necessary
        
        # Run LibreOffice in headless mode to convert DOCX to PDF
        subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', docx_file_path, '--outdir', output_dir], check=True)
    
    elif current_os == 'Linux':
        # On Linux/Ubuntu, use LibreOffice (assumes it's in the PATH)
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', docx_file_path, '--outdir', output_dir], check=True)
    
    else:
        raise RuntimeError(f"Unsupported operating system: {current_os}")
    
    # Check if the PDF was created
    if not os.path.exists(pdf_file_path):
        raise RuntimeError(f"Failed to convert {docx_file_path} to PDF.")
    
    return pdf_file_path
